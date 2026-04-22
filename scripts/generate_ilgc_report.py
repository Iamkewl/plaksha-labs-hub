from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "HEXA25-ILGC-ALL STUFF-SEM2"
ASSET_DIR = SOURCE_DIR / "report-assets"
OUTPUT_PATH = SOURCE_DIR / "Hexa25_ILGC_Project_Report.docx"

ACCENT = RGBColor(0x0B, 0x7D, 0x77)
ACCENT_HEX = "#0B7D77"
ACCENT_DARK = "#0A4F4A"
ACCENT_LIGHT = "#DDF4F2"
TEXT_DARK = "#1F2937"
MUTED = "#6B7280"
HIGHLIGHT = "#F2A23A"
ALERT = "#D65745"

REPORT_TITLE = "Plaksha Makerspace Hub"
REPORT_SUBTITLE = (
    "ILGC-2 Report: Solving the Makerspace Accessibility and Information Gap "
    "through an integrated operational platform"
)
REPORT_DATE = "22 April 2026"

TEAM_MEMBERS = [
    {
        "name": "Dhiraj Deva",
        "role": "Problem framing and coordination-flow ideation",
        "evidence": "Week_11 brainstorming board: MS Forms and chat-based booking alternatives.",
    },
    {
        "name": "Mauryan Jaiswal",
        "role": "Inventory automation and operations logic",
        "evidence": "Week_11 brainstorming board: chart system and automated inventory shelf concepts.",
    },
    {
        "name": "Suryaansh",
        "role": "Core web-product architecture and training enablement",
        "evidence": "Week_11 brainstorming board: web app, machine status, and training-module directions.",
    },
    {
        "name": "Greeshma Tanvi Reddy Tadi",
        "role": "Access verification and QR-driven interaction design",
        "evidence": "Week_11 brainstorming board: face-scan verification and QR inventory ideas.",
    },
    {
        "name": "Priyanka Chopurala",
        "role": "Service-design alternatives and user-side workflow exploration",
        "evidence": "Week_11 brainstorming board: physical logbook, slot-slip, and chatbot alternatives.",
    },
    {
        "name": "Harshil Jain",
        "role": "Structured data and booking-system thinking",
        "evidence": "Week_11 brainstorming board: materials database and ID-based booking concepts.",
    },
]

PRIORITY_SCORES = [
    ("Professors repository", 15),
    ("Lost-and-found system", 13),
    ("Makerspace info gap", 12),
    ("Inclusive infrastructure", 9),
]

SOLUTION_CLUSTERS = [
    ("Centralized visibility tools", 5),
    ("Low-tech coordination", 4),
    ("Automation and verification", 3),
    ("Training and guidance", 2),
]

TIMELINE_STEPS = [
    ("Week 6", "Zone scanning and SWOT shortlist"),
    ("Week 11", "Brainstorming, dot voting, and solution clustering"),
    ("Convergence", "Shift from fragmented tools to a unified hub"),
    ("Prototype", "Operational modules implemented in the platform"),
    ("Next phase", "Public discovery, BOM approvals, and dual-support booking"),
]

MODULE_COVERAGE = [
    ("Auth and access control", 1.0),
    ("Catalogs", 1.0),
    ("Bookings", 1.0),
    ("Projects and BOM", 1.0),
    ("Material approvals", 0.9),
    ("Mentor availability", 0.95),
    ("Training management", 0.95),
    ("Notifications", 0.95),
    ("Analytics", 0.9),
    ("Purchase orders", 0.35),
]

RESOURCE_ROWS = [
    ("Laptop and internet access", "Existing team development setup", 0),
    ("Hosting and deployment", "Prototype-grade hosting on free/student tier", 0),
    ("Database and auth services", "Prototype tier for pilot validation", 0),
    ("QR labels and signage", "Machine and material navigation aids", 1500),
    ("Onboarding posters and guides", "Printed quick-start and safety guidance", 1000),
    ("Domain, email, and ops reserve", "Pilot-ready deployment buffer", 2000),
    ("Contingency", "Unexpected setup and testing requirements", 2000),
]

AI_TOOL_ROWS = [
    (
        "ChatGPT (GPT-5.3)",
        "Brainstorming problem causes, presentation structuring, and solution comparison.",
        "The team filtered outputs against on-ground observation, stakeholder context, and campus feasibility before accepting any recommendation.",
    ),
    (
        "GitHub Copilot (GPT-5.4)",
        "Prototype implementation support, codebase summarization, chart generation, and Word-report assembly.",
        "All facts were constrained to the uploaded project material and the live codebase; the final report structure and wording were human-reviewed and edited.",
    ),
]

SOURCES = [
    ("[S1]", "Group 25- SWOT.pptx, Week 6 ideation and prioritization deck."),
    ("[S2]", "Week_11.pptx, brainstorming, dot-voting, and shortlisted solution concepts."),
    ("[S3]", "ILGC Presentation Content.docx, synthesized stakeholder and jury-preparation notes."),
    ("[S4]", "Project Report Template.docx, required report structure and submission sections."),
    ("[S5]", "state.md, current implemented product scope and feature inventory."),
    ("[S6]", "roadmap.md, planned next-phase product refinements and implementation logic."),
    ("[S7]", "chats/main-page-redesign-samples/01-foundry-executive.html, selected Concept A homepage prototype."),
    ("[S8]", "agent-control/.agents/knowledge/Tasks/task-2026-04-21-concept-a-rollout.md, evidence of Concept A selection and rollout."),
    ("[S9]", "src/app/page.tsx, current public homepage implementation derived from the selected concept."),
    ("[S10]", "prisma/schema.prisma, current domain entities and operational data model."),
]

SWOT_ITEMS = {
    "Strengths": [
        "Improves resource visibility and utilization, reducing confusion and inefficiencies.",
        "Low-to-moderate technology implementation cost for a pilot-stage solution.",
        "Encourages innovation and student engagement by making the makerspace easier to approach.",
    ],
    "Weaknesses": [
        "Needs regular data updates to remain trustworthy.",
        "Requires staff onboarding and disciplined digital ownership.",
        "Booking-system integration and approval-state coordination increase product complexity.",
    ],
    "Opportunities": [
        "Can include tutorials, onboarding, and project showcases for beginners.",
        "Can become a collaboration hub for student clubs, teams, and startups.",
        "Can produce analytics that justify funding, staffing, and makerspace expansion.",
    ],
    "Threats": [
        "Outdated data would quickly erode trust in the platform.",
        "Low awareness could preserve underutilization even after launch.",
        "The platform could stall without clear administrative ownership and upkeep.",
    ],
}

STAKEHOLDER_FEATURE_ROWS = [
    (
        "Students planning a visit",
        "Need to know what exists, what is available, and what they are allowed to use before reaching the makerspace.",
        "Searchable catalogs, machine status, material visibility, public projects, and a clear sign-in path.",
    ),
    (
        "Beginners and first-time users",
        "Need guidance, training cues, and a lower-friction way to understand how makerspace processes work.",
        "Training gates, mentor-linked workflows, safety/context information, and onboarding-oriented content.",
    ),
    (
        "Mentors and faculty",
        "Need predictable coordination around availability, support sessions, and supervised project execution.",
        "Availability management, mentor session booking, supervised project views, and approval handoffs.",
    ),
    (
        "Admins and operations owners",
        "Need inventory governance, approval visibility, usage analytics, and low-stock awareness.",
        "Admin dashboards, request queues, analytics, notifications, and role-protected management pages.",
    ),
]

CONCEPT_COMPARISON_ROWS = [
    (
        "Concept A - Foundry Executive",
        "Light command-console hierarchy with KPI-first structure and operational scanning.",
        "Strong fit for makerspace workflows, role clarity, and system trust.",
        "Can feel too dashboard-sensitive if copied directly onto public pages.",
        "Selected and rolled out as the base visual/structural direction [S7][S8].",
    ),
    (
        "Concept B - Parchment and Steel",
        "Editorial, archival, operations-board aesthetic with strong narrative rhythm.",
        "Excellent for storytelling and institutional tone.",
        "Less efficient for dense operational actions and fast scanning.",
        "Useful reference for content framing, but not chosen as the primary system language.",
    ),
    (
        "Concept C - Vanguard Runway",
        "High-tempo, split-layout hierarchy with strong motion and visual energy.",
        "Memorable and striking for promotional presentation.",
        "More presentation-led than governance-led for a makerspace operations system.",
        "Rejected as the core direction because operational clarity mattered more than spectacle.",
    ),
]

CONCEPT_TO_FINAL_ROWS = [
    (
        "Public homepage structure",
        "Concept A used a KPI-heavy operations preview with filters, inventory table, alerts, and quick actions.",
        "The live homepage keeps the command-surface logic but shifts to a narrative hero, role journey, and public-project discovery.",
        "This preserves operational credibility while preventing public visitors from seeing internal dashboard-style content [S7][S8][S9].",
    ),
    (
        "Primary action pattern",
        "Concept A centered on direct workspace entry and operational scanning.",
        "The final page keeps Sign In as the main action but adds a Browse Public Projects path for unauthenticated users.",
        "The system serves both internal operators and curious public visitors, so the public route must support both intents.",
    ),
    (
        "Information density",
        "The original concept presented machine/material data directly on the homepage.",
        "The final implementation moves operational detail into authenticated routes and uses concise summaries on public surfaces.",
        "This is safer, clearer, and more appropriate for a campus-facing front door.",
    ),
    (
        "Visual tone",
        "Concept A was light, crisp, and command-oriented.",
        "The final implementation translates that hierarchy into a more cinematic, darker brand language used across the project.",
        "The rollout kept the operational grammar of Concept A while adapting it to a stronger institutional identity [S8][S9].",
    ),
]

DESIGN_BRIEF_ROWS = [
    (
        "Core problem",
        "Students, mentors, and admins lack one reliable system for discovery, access, approvals, and operational tracking in the makerspace.",
    ),
    (
        "Primary users",
        "Students, beginner makers, mentors/faculty, and admin operators.",
    ),
    (
        "Primary objective",
        "Create a single operational layer that reduces uncertainty before a visit and improves coordination during project execution.",
    ),
    (
        "Desired behavior change",
        "Move users from ad-hoc, channel-by-channel coordination to a structured discover-book-build-govern workflow.",
    ),
    (
        "Success condition",
        "Users can discover resources, understand eligibility, book support, manage projects/BOMs, and receive clear operational feedback without manual confusion.",
    ),
    (
        "Non-negotiables",
        "Role-aware access control, live operational status, beginner guidance, project accountability, and public/private separation.",
    ),
    (
        "Delivery constraints",
        "The solution must remain software-first, pilot-feasible, responsive, and compatible with the current Next.js + Prisma stack.",
    ),
]

DESIGN_SPEC_ROWS = [
    (
        "Information visibility",
        "Machine and material records must expose key operating data such as status, location, training needs, stock, and threshold.",
        "The problem begins with uncertainty before arrival.",
        "Catalog routes and detail pages in the current prototype [S5].",
    ),
    (
        "Guided access",
        "Booking must enforce future-only scheduling, machine availability, overlap checks, and training requirements.",
        "Access without rules recreates friction in a different form.",
        "Booking flows and action-layer constraints [S5].",
    ),
    (
        "Mentor-supported workflows",
        "Support-required machines must be able to reserve machine plus mentor together.",
        "Some tools cannot be treated as self-serve resources.",
        "Roadmap direction and schema support flag [S6][S10].",
    ),
    (
        "Project accountability",
        "Projects must support members, mentors, BOM versions, and approval-linked request history.",
        "The makerspace is project-driven, not only transaction-driven.",
        "Project/BOM scope in the current prototype [S5].",
    ),
    (
        "Governance and permissions",
        "Admin and mentor actions must be separated through role-aware routing and server-side permission checks.",
        "Operational trust depends on controlled actions.",
        "Middleware, guards, and admin/mentor routes [S5].",
    ),
    (
        "Public storytelling",
        "Public users must be able to understand the makerspace and browse public projects without accessing sensitive operational data.",
        "The platform should also act as an engagement surface.",
        "Homepage and public-project direction [S6][S8][S9].",
    ),
    (
        "Accessibility and responsiveness",
        "The interface should preserve visible focus states, keyboard clarity, responsive layout behavior, and reduced-motion safety.",
        "The audience includes different user types and devices.",
        "Concept and implementation patterns in the chosen public route [S7][S9].",
    ),
    (
        "Operational resilience",
        "Public pages must fail gracefully when live project data is unavailable.",
        "A makerspace front door cannot disappear during infrastructure issues.",
        "Homepage fallback pattern in the live implementation [S9].",
    ),
]

SCREEN_BRIEF_ROWS = [
    (
        "Public homepage",
        "Visitors and prospective users",
        "Explain the platform, clarify role journeys, and surface public projects.",
        "Narrative hero, operational flow card, role cards, and public showcase entry points.",
    ),
    (
        "Sign-in and auth entry",
        "Students, mentors, admins",
        "Move users into the correct protected workspace with minimum confusion.",
        "Clear brand continuity, strong primary action, and role-safe access handling.",
    ),
    (
        "Dashboard",
        "Authenticated users by role",
        "Summarize the most relevant tasks, risks, and actions for each user type.",
        "Different KPI and workflow emphasis for students, mentors, and admins.",
    ),
    (
        "Machine and material catalogs",
        "Students and admins",
        "Support remote planning through search, filtering, and detail visibility.",
        "Status, location, thresholds, specs, and training clues presented before action.",
    ),
    (
        "Booking flow",
        "Students and mentors",
        "Convert intent into valid reservations with policy-aware checks.",
        "Dual-tab booking, overlap detection, training checks, and mentor linkage where required.",
    ),
    (
        "Project and BOM flows",
        "Student teams, mentors, admins",
        "Connect project execution with resource planning and approvals.",
        "Member roles, BOM versioning, request lifecycle, and notification events.",
    ),
    (
        "Admin operations pages",
        "Admins",
        "Oversee users, machines, materials, requests, training, and analytics from one control layer.",
        "Structured review queues and management surfaces instead of scattered manual processes.",
    ),
]

TECH_STACK_ROWS = [
    ("Presentation layer", "Next.js App Router with React and TypeScript", "Supports route-based public and authenticated experiences in one codebase."),
    ("UI primitives", "Custom components in src/components/ui", "Provides reusable cards, buttons, tables, tabs, forms, and layout surfaces."),
    ("Server logic", "Server actions in src/app/actions", "Holds workflow rules for bookings, projects, BOMs, approvals, notifications, and analytics."),
    ("Data model", "Prisma with PostgreSQL", "Stores users, machines, materials, bookings, projects, BOMs, requests, training, and notifications."),
    ("Authentication", "NextAuth with Prisma adapter", "Enables managed sessions and role-aware access."),
    ("Permissions and resilience", "Middleware, auth guards, and rate limiting", "Protects routes, enforces roles, and reduces abuse on high-frequency endpoints."),
    ("Communication", "In-app notifications with optional email provider path", "Supports operational awareness and future transactional communications."),
]

ENTITY_ROWS = [
    ("User", "Identity, role, and system participation", "Connects to bookings, trainings, projects, availability, approvals, and notifications."),
    ("Machine", "Tool inventory and access constraints", "Carries status, location, pricing, specifications, safety, and mentor/training requirements."),
    ("Material", "Consumable inventory and stock governance", "Tracks quantity, threshold, pricing, and issue patterns."),
    ("Booking", "Time-bound reservation record", "Links users to machines and/or mentors with status and conflict rules."),
    ("Training", "Capability gate for machine use", "Validates whether a user is allowed to book a training-required machine."),
    ("Project", "Team-level making activity", "Holds mentor assignment, members, BOM history, and visibility intent."),
    ("ProjectMember", "Role inside a project", "Separates leads from members for permission-sensitive actions."),
    ("Bom and BomItem", "Structured material planning", "Versioned project resource plan that can flow into approvals and requests."),
    ("MaterialRequest", "Operational request and issue lifecycle", "Tracks approval, issuance, and requester/approver metadata."),
    ("MentorAvailability", "Support-slot publication", "Allows mentors to expose recurring or one-off time windows."),
    ("Notification", "Workflow awareness record", "Delivers state changes back to the relevant actors."),
]

WORKFLOW_ROWS = [
    (
        "Machine discovery to booking",
        "Student explores catalog, opens details, checks eligibility, and books an available slot.",
        "Availability, future date, overlap, and training validation.",
        "Confirmed or pending booking plus notification-ready audit trail.",
    ),
    (
        "Mentor-supported machine booking",
        "Student chooses a machine that requires support and reserves both machine and mentor.",
        "Machine + mentor conflict checks and mandatory mentor selection.",
        "A single governed reservation for advanced or supervised tool use.",
    ),
    (
        "Project to BOM to material request",
        "Team lead creates a project, assembles BOM items, submits for approval, and triggers request handling.",
        "Version control, approver identity, admin gating, and quantity/stock logic.",
        "Trackable resource planning instead of ad-hoc asking.",
    ),
    (
        "Public project storytelling",
        "Public visitor browses visible projects without entering protected operational areas.",
        "Visibility controls and safe public data boundaries.",
        "Campus-facing discovery and engagement layer for makerspace work.",
    ),
]

VALIDATION_ROWS = [
    (
        "Functional workflow validation",
        "Confirm route and action coverage for auth, booking, projects, BOMs, materials, notifications, and analytics.",
        "Documented implemented feature inventory [S5].",
        "Run representative student, mentor, and admin task paths end to end during pilot.",
    ),
    (
        "Permission and safety validation",
        "Verify that public, mentor, and admin surfaces remain correctly isolated.",
        "Role-aware route protection and auth-guard logic documented in project state [S5].",
        "Attempt unauthorized access paths and confirm they fail safely.",
    ),
    (
        "Data-model validation",
        "Check that entities, statuses, and relationships support the planned workflows.",
        "Current schema and roadmap expansion notes [S6][S10].",
        "Test lifecycle edges such as partial issuance, BOM approval, and mentor-linked booking.",
    ),
    (
        "Design validation",
        "Measure whether the selected design language preserves clarity on public and authenticated surfaces.",
        "Concept A selection and rollout evidence [S7][S8][S9].",
        "Ask users to complete discovery and sign-in tasks while recording confusion points and completion time.",
    ),
    (
        "Pilot success metrics",
        "Track operational improvement after deployment.",
        "Current prototype provides the required modules but not yet a full pilot dataset.",
        "Measure discovery time, booking completion, approval turnaround, low-stock awareness, and project-publication engagement.",
    ),
]

ROADMAP_PHASE_ROWS = [
    (
        "Phase 1",
        "Data model and validation foundation",
        "New material-request statuses, BOM approval metadata, mentor-support flag, and validation updates.",
        "Schema compiles cleanly and new constraints map safely to existing workflows.",
    ),
    (
        "Phase 2",
        "BOM allocation workflow",
        "Student/mentor request submission, BOM approver review, and admin-stage gating.",
        "Requests cannot bypass the BOM approval stage before admin action.",
    ),
    (
        "Phase 3",
        "Mentor-supported machine booking",
        "Machine flagging, mentor requirement enforcement, dual conflict checks, and single reservation persistence.",
        "Support-required bookings fail safely when no mentor is selected.",
    ),
    (
        "Phase 4",
        "Public project discovery",
        "Visibility controls, public projects route, and landing-page integration.",
        "Visitors can browse public projects with no auth regression.",
    ),
    (
        "Phase 5",
        "QA and stabilization",
        "Permission checks, workflow regression review, and state-document updates.",
        "No unauthorized bypass and no regressions in existing core workflows.",
    ),
]


def format_inr(amount: int) -> str:
    return f"Rs {amount:,.0f}"


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update field in Word"
    separate.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def shade_cell(cell, fill: str):
    props = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    props.append(shade)


def style_document(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = ACCENT

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0x22, 0x35, 0x46)

    heading3 = document.styles["Heading 3"]
    heading3.font.name = "Calibri"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = ACCENT


def add_header_footer(document: Document):
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header.add_run("Hexa25 ILGC Report")
        header_run.font.name = "Calibri"
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Page ")
        footer_run.font.name = "Calibri"
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        add_field(footer, "PAGE")


def add_title_page(document: Document, banner_path: Path):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(banner_path), width=Inches(6.8))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_TITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(REPORT_SUBTITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    document.add_paragraph()

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Project Title: Plaksha Makerspace Hub\n")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run = meta.add_run("Team Number: 25\n")
    run.font.size = Pt(11)
    run = meta.add_run("Date: " + REPORT_DATE)
    run.font.size = Pt(11)

    members = document.add_paragraph()
    members.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = members.add_run("Team Members and Roles\n")
    run.font.bold = True
    run.font.size = Pt(12)
    for member in TEAM_MEMBERS:
        line = members.add_run(f"{member['name']} - {member['role']}\n")
        line.font.size = Pt(10.5)
        line.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Prepared from the shared HEXA25 evidence pack and the current working prototype repository."
    )
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_page_break()


def add_toc_page(document: Document):
    document.add_heading("Contents", level=1)
    paragraph = document.add_paragraph()
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    note = document.add_paragraph()
    note_run = note.add_run("If the table of contents does not populate automatically, open the file in Word and update fields once.")
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    document.add_page_break()


def add_figure(document: Document, image_path: Path, caption: str, width_inches: float = 6.3):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_quote(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.right_indent = Cm(0.4)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_table(document: Document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    head_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = head_cells[index]
        cell.text = header
        shade_cell(cell, "DDF4F2")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = ACCENT
                run.font.size = Pt(10.5)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    document.add_paragraph()


def build_cover_banner(path: Path):
    fig, ax = plt.subplots(figsize=(11, 3.2), dpi=180)
    fig.patch.set_facecolor(ACCENT_DARK)
    ax.set_facecolor(ACCENT_DARK)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    ax.add_patch(FancyBboxPatch((0.03, 0.2), 0.17, 0.6, boxstyle="round,pad=0.02,rounding_size=0.03", linewidth=0, facecolor="#0E625C"))
    ax.add_patch(FancyBboxPatch((0.23, 0.28), 0.1, 0.44, boxstyle="round,pad=0.02,rounding_size=0.03", linewidth=0, facecolor="#15807B"))
    ax.add_patch(FancyBboxPatch((0.35, 0.14), 0.11, 0.72, boxstyle="round,pad=0.02,rounding_size=0.03", linewidth=0, facecolor="#0B9A91"))

    ax.text(0.52, 0.68, "HEXA25", fontsize=26, fontweight="bold", color="white", ha="left", va="center")
    ax.text(0.52, 0.47, "Makerspace Accessibility and Information Gap", fontsize=16, color="#D9FFFA", ha="left", va="center")
    ax.text(0.52, 0.29, "Integrated platform report for ILGC-2", fontsize=12.5, color="#BDEBE6", ha="left", va="center")

    fig.savefig(path, bbox_inches="tight", pad_inches=0.1)
    plt.close(fig)


def build_priority_chart(path: Path):
    labels = [item[0] for item in PRIORITY_SCORES][::-1]
    values = [item[1] for item in PRIORITY_SCORES][::-1]
    colors = ["#CBD5E1", ALERT, HIGHLIGHT, "#A7F3D0"]

    fig, ax = plt.subplots(figsize=(8.2, 4.6), dpi=180)
    bars = ax.barh(labels, values, color=colors, edgecolor="#1F2937")
    ax.set_title("Opportunity Prioritization Across Campus Problems", fontsize=14, pad=12)
    ax.set_xlabel("Composite score (ease + feasibility + cost)", fontsize=10)
    ax.set_xlim(0, 16)
    ax.grid(axis="x", alpha=0.25)
    ax.spines[["top", "right"]].set_visible(False)
    for bar, value in zip(bars, values):
        ax.text(value + 0.2, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=10, color="#111827")
    ax.text(7.6, 0.2, "Chosen challenge based on score plus strategic fit with makerspace impact and digital prototyping feasibility.", fontsize=8.8, color=MUTED, ha="center")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_solution_cluster_chart(path: Path):
    labels = [item[0] for item in SOLUTION_CLUSTERS]
    values = [item[1] for item in SOLUTION_CLUSTERS]
    colors = [ACCENT_HEX, HIGHLIGHT, "#7C3AED", "#F97316"]

    fig, ax = plt.subplots(figsize=(7.8, 4.6), dpi=180)
    ax.bar(labels, values, color=colors)
    ax.set_title("Idea Clusters Emerging from Team Brainstorming", fontsize=14, pad=12)
    ax.set_ylabel("Number of ideas", fontsize=10)
    ax.grid(axis="y", alpha=0.25)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(axis="x", rotation=15)
    for index, value in enumerate(values):
        ax.text(index, value + 0.08, str(value), ha="center", fontsize=10, color="#111827")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_timeline_chart(path: Path):
    fig, ax = plt.subplots(figsize=(9.2, 2.8), dpi=180)
    ax.set_xlim(0, len(TIMELINE_STEPS) - 1)
    ax.set_ylim(-1.4, 1.4)
    ax.axis("off")
    ax.plot([0, len(TIMELINE_STEPS) - 1], [0, 0], color="#94A3B8", linewidth=2)
    for index, (title, description) in enumerate(TIMELINE_STEPS):
        ax.scatter(index, 0, s=180, color=ACCENT_HEX, zorder=3)
        ax.text(index, 0.42 if index % 2 == 0 else -0.62, title, ha="center", va="center", fontsize=10, fontweight="bold", color="#0F172A")
        ax.text(index, 0.82 if index % 2 == 0 else -1.02, description, ha="center", va="center", fontsize=8.8, color=MUTED, wrap=True)
    ax.set_title("Refinement Timeline from Problem Selection to Prototype Direction", fontsize=14, pad=8)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_architecture_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(9.4, 5.4), dpi=180)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")

    def box(x, y, w, h, text, fill, text_color="#111827"):
        patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.18", facecolor=fill, edgecolor="#1F2937", linewidth=1.2)
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=10, color=text_color, wrap=True)

    box(0.4, 4.8, 2.0, 1.0, "Students", "#E2F5F3")
    box(0.4, 3.3, 2.0, 1.0, "Mentors / Faculty", "#E2F5F3")
    box(0.4, 1.8, 2.0, 1.0, "Admins", "#E2F5F3")

    box(3.2, 5.0, 2.8, 0.9, "Discovery Layer\nMachine and material catalog", "#DDF4F2")
    box(3.2, 3.8, 2.8, 0.9, "Execution Layer\nBooking, training, mentor support", "#DDF4F2")
    box(3.2, 2.6, 2.8, 0.9, "Project Layer\nProjects, BOMs, approvals", "#DDF4F2")
    box(3.2, 1.4, 2.8, 0.9, "Operations Layer\nNotifications and analytics", "#DDF4F2")

    box(6.8, 4.1, 2.3, 1.0, "Prisma + PostgreSQL\nOperational data store", "#FEF3C7")
    box(6.8, 2.3, 2.3, 1.0, "Governance Controls\nRoles, rate limits, audit trail", "#FEF3C7")

    for y in [5.25, 3.75, 2.25]:
        ax.annotate("", xy=(3.2, y), xytext=(2.4, y), arrowprops=dict(arrowstyle="-|>", color="#475569", linewidth=1.5))
    for y in [5.45, 4.25, 3.05, 1.85]:
        ax.annotate("", xy=(6.8, y), xytext=(6.0, y), arrowprops=dict(arrowstyle="-|>", color="#475569", linewidth=1.5))

    ax.set_title("High-Level Proposed System Architecture", fontsize=14, pad=10)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_module_coverage_chart(path: Path):
    labels = [item[0] for item in MODULE_COVERAGE][::-1]
    values = [item[1] for item in MODULE_COVERAGE][::-1]
    colors = [ACCENT_HEX if value >= 0.9 else HIGHLIGHT if value >= 0.7 else ALERT for value in values]

    fig, ax = plt.subplots(figsize=(8.2, 4.8), dpi=180)
    bars = ax.barh(labels, values, color=colors)
    ax.set_xlim(0, 1.05)
    ax.set_xlabel("Current prototype readiness", fontsize=10)
    ax.set_title("Current Prototype Coverage by Subsystem", fontsize=14, pad=12)
    ax.grid(axis="x", alpha=0.25)
    ax.spines[["top", "right"]].set_visible(False)
    for bar, value in zip(bars, values):
        ax.text(value + 0.02, bar.get_y() + bar.get_height() / 2, f"{int(value * 100)}%", va="center", fontsize=9.5, color="#111827")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_swot_matrix(path: Path):
    fig, ax = plt.subplots(figsize=(9.4, 6.2), dpi=180)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")

    quadrants = [
        (0.4, 5.2, 4.35, 4.1, "Strengths", SWOT_ITEMS["Strengths"], "#E2F5F3"),
        (5.2, 5.2, 4.35, 4.1, "Weaknesses", SWOT_ITEMS["Weaknesses"], "#FFF1D9"),
        (0.4, 0.6, 4.35, 4.1, "Opportunities", SWOT_ITEMS["Opportunities"], "#E7F0FF"),
        (5.2, 0.6, 4.35, 4.1, "Threats", SWOT_ITEMS["Threats"], "#FDE7E5"),
    ]

    for x, y, width, height, title, items, fill in quadrants:
        patch = FancyBboxPatch(
            (x, y),
            width,
            height,
            boxstyle="round,pad=0.05,rounding_size=0.18",
            facecolor=fill,
            edgecolor="#1F2937",
            linewidth=1.1,
        )
        ax.add_patch(patch)
        ax.text(x + 0.22, y + height - 0.45, title, fontsize=12, fontweight="bold", color="#0F172A", ha="left", va="center")
        body = "\n".join([f"• {item}" for item in items])
        ax.text(x + 0.22, y + height - 0.82, body, fontsize=8.7, color="#374151", ha="left", va="top", wrap=True)

    ax.set_title("SWOT Analysis of the Selected Makerspace Platform Direction", fontsize=14, pad=10)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_workflow_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(10.2, 3.6), dpi=180)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")

    steps = [
        (0.5, "Discover\nresources"),
        (2.6, "Check eligibility\nand support needs"),
        (4.9, "Book or create\nproject/BOM request"),
        (7.3, "Mentor / approver\nreview gate"),
        (9.4, "Admin action,\nissue, or confirm"),
        (11.2, "Notify and\nmeasure"),
    ]

    for index, (x, label) in enumerate(steps):
        patch = FancyBboxPatch(
            (x, 1.25),
            1.45,
            1.2,
            boxstyle="round,pad=0.04,rounding_size=0.16",
            facecolor="#E2F5F3" if index % 2 == 0 else "#FEF3C7",
            edgecolor="#1F2937",
            linewidth=1.1,
        )
        ax.add_patch(patch)
        ax.text(x + 0.725, 1.85, label, ha="center", va="center", fontsize=9.2, color="#111827")
        if index < len(steps) - 1:
            ax.annotate(
                "",
                xy=(x + 1.65, 1.85),
                xytext=(steps[index + 1][0] - 0.18, 1.85),
                arrowprops=dict(arrowstyle="-|>", color="#475569", linewidth=1.4),
            )

    ax.text(6, 3.1, "Operational workflow from discovery to governance", ha="center", fontsize=14, fontweight="bold", color="#0F172A")
    ax.text(6, 0.6, "The system is designed to move users through visibility, validation, execution, review, and feedback without channel switching.", ha="center", fontsize=9, color="#6B7280")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "banner": ASSET_DIR / "cover-banner.png",
        "priority": ASSET_DIR / "priority-matrix.png",
        "clusters": ASSET_DIR / "solution-clusters.png",
        "swot": ASSET_DIR / "swot-matrix.png",
        "timeline": ASSET_DIR / "refinement-timeline.png",
        "architecture": ASSET_DIR / "system-architecture.png",
        "workflow": ASSET_DIR / "workflow-diagram.png",
        "coverage": ASSET_DIR / "module-coverage.png",
    }
    build_cover_banner(assets["banner"])
    build_priority_chart(assets["priority"])
    build_solution_cluster_chart(assets["clusters"])
    build_swot_matrix(assets["swot"])
    build_timeline_chart(assets["timeline"])
    build_architecture_diagram(assets["architecture"])
    build_workflow_diagram(assets["workflow"])
    build_module_coverage_chart(assets["coverage"])
    return assets


def add_cover_note(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.3)
    paragraph.paragraph_format.right_indent = Cm(0.3)
    run = paragraph.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)


def save_document(document: Document, path: Path) -> Path:
    try:
        document.save(path)
        return path
    except PermissionError:
        fallback_path = path.with_name(f"{path.stem}_{datetime.now():%Y%m%d_%H%M%S}{path.suffix}")
        document.save(fallback_path)
        return fallback_path


def build_report():
    assets = build_assets()
    document = Document()
    style_document(document)
    add_header_footer(document)
    add_title_page(document, assets["banner"])
    add_toc_page(document)

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "Hexa25's project responds to a recurring operational problem at Plaksha Makerspace: students, mentors, and staff do not have a single reliable system to discover tools, check machine availability, understand training requirements, view material stock, and coordinate project approvals. The shared evidence pack repeatedly points to fragmented access, repeated uncertainty, and lost time as the dominant user experience."  # noqa: E501
    )
    document.add_paragraph(
        "The team's answer is an integrated digital hub rather than a single-purpose fix. The proposed platform combines machine and material discovery, mentor-supported booking, project and BOM workflows, approval stages, notifications, and analytics in one operational layer. This allows the makerspace to function as a coordinated campus resource rather than a set of disconnected manual interactions."  # noqa: E501
    )
    document.add_paragraph(
        "This direction is already feasible. The current prototype scope confirms implemented modules for authentication, catalogs, bookings, project management, BOM handling, material request flows, mentor availability, training records, notifications, and analytics. The report therefore argues not only that the solution is desirable, but also that it is buildable, scalable, and institutionally useful."  # noqa: E501
    )
    document.add_paragraph("Expected outcome:", style="Heading 3")
    for item in [
        "Reduce student uncertainty before a makerspace visit by making access conditions visible in advance.",
        "Lower coordination load on mentors and admins through structured booking, approval, and notification workflows.",
        "Improve utilization and planning through analytics on machines, bookings, material demand, and user activity.",
        "Create a platform that can grow from a prototype into a campus operations layer with public showcase potential.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("2. Background Research", level=1)
    document.add_heading("2.1 Problem Identification and Opportunity Framing", level=2)
    document.add_paragraph(
        "The Week 6 team deck shows that the group initially scanned several campus opportunities across different zones before converging on the makerspace challenge. A weighted priority matrix compared lost-and-found management, campus accessibility gaps, a professors' repository, and the makerspace information gap against ease of implementation, feasibility, and cost [S1]."  # noqa: E501
    )
    document.add_paragraph(
        "Although the professors' repository problem scored highest numerically, the makerspace challenge emerged as the most strategically aligned opportunity. It sat at the intersection of strong campus relevance, clear student pain, and the team's ability to deliver an actionable digital prototype. In other words, the final choice was driven by score plus fit, not score alone."  # noqa: E501
    )
    add_figure(
        document,
        assets["priority"],
        "Figure 1. Opportunity prioritization across shortlisted campus problems, synthesized from the Week 6 SWOT deck [S1].",
    )
    document.add_heading("2.2 Existing Solutions and What They Miss", level=2)
    add_table(
        document,
        ["Existing approach", "What it can do", "Why it is insufficient"],
        [
            ("WhatsApp / Excel / MS Forms", "Accepts lightweight requests and manual coordination.", "No live inventory or machine state, weak ownership, and poor auditability."),
            ("Physical logbooks and slips", "Creates a basic record at the point of use.", "Not discoverable remotely, error-prone, and hard to analyze over time."),
            ("Standalone status chart", "Improves local visibility of machine availability.", "Does not connect bookings, training, or materials in one workflow."),
            ("QR or ID access-only systems", "Helps authenticate or verify access.", "Solves entry friction but not planning, guidance, or approvals."),
            ("Sensor-heavy smart shelves", "Could automate inventory tracking.", "Higher hardware complexity than needed for a pilot and does not solve coordination alone."),
        ],
    )
    document.add_paragraph(
        "The team's own ideation makes the gap clear: almost every alternative addressed only one slice of the problem. Some ideas improved discoverability, some improved verification, and some improved automation, but none created a full makerspace workflow from awareness to booking to material use to reporting. The background research therefore pushed the team away from fragmented fixes and toward an integrated platform strategy."  # noqa: E501
    )
    document.add_heading("2.3 Gap Analysis", level=2)
    for item in [
        "Fragmentation: existing workarounds solve isolated tasks rather than the full user journey.",
        "Low visibility: students cannot reliably know machine status, material stock, or support availability before arriving.",
        "Weak onboarding: beginners need guidance, tutorials, and training clarity, not only access control.",
        "No operational memory: manual methods produce weak history, low accountability, and limited data for planning.",
        "Poor scalability: once makerspace usage increases, disconnected channels become slower and less trustworthy.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    add_figure(
        document,
        assets["clusters"],
        "Figure 2. Solution clusters that emerged from Week 11 brainstorming, showing how early ideas spread across visibility, coordination, automation, and guidance [S2].",
    )
    document.add_heading("2.4 SWOT Analysis of the Selected Direction", level=2)
    document.add_paragraph(
        "The original SWOT exercise for the makerspace challenge is strong because it does not describe the idea as universally positive. It correctly identifies that a centralized platform can create clarity and engagement, but only if the data stays current and someone owns the operating discipline behind it [S1]. That makes SWOT especially important here: the platform is not merely a design artifact, it is an operational system whose value rises or collapses based on upkeep."  # noqa: E501
    )
    add_figure(
        document,
        assets["swot"],
        "Figure 3. SWOT matrix for the selected makerspace platform direction, synthesized from the Week 6 SWOT deck and interpreted against the current prototype [S1][S5].",
    )
    add_table(
        document,
        ["Strengths", "Weaknesses", "Opportunities", "Threats"],
        [(
            "\n".join(SWOT_ITEMS["Strengths"]),
            "\n".join(SWOT_ITEMS["Weaknesses"]),
            "\n".join(SWOT_ITEMS["Opportunities"]),
            "\n".join(SWOT_ITEMS["Threats"]),
        )],
    )

    document.add_heading("3. Stakeholder Feedback on the Proposed Solution", level=1)
    document.add_paragraph(
        "The local evidence pack documents direct observation of makerspace use, repeated site visits, informal interaction with students, admins, and faculty, and a presentation script prepared around synthesized survey and interview findings [S3]. The raw survey spreadsheet was not preserved in the uploaded folder, but the available documents still capture the dominant stakeholder concerns with enough consistency to guide design."  # noqa: E501
    )
    document.add_heading("3.1 Methods Used", level=2)
    for item in [
        "Repeated makerspace visits to observe how students search for tools, plan work, and request access.",
        "Informal interviews and contextual conversations with students, admins, and faculty mentioned in the presentation content [S3].",
        "Team-level synthesis through SWOT analysis, dot voting, and problem reframing [S1][S2].",
        "Prototype-grounded review using the current implemented platform scope and roadmap [S5][S6].",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("3.2 Stakeholder Needs Summary", level=2)
    add_table(
        document,
        ["Stakeholder", "Documented need", "Design implication"],
        [
            ("Students", "Need to know what exists, what is available, and how to access it before arriving.", "Prioritize searchable catalogs, machine status, and beginner-friendly guidance."),
            ("Beginners", "Need support before independent use of machines and materials.", "Expose training requirements, mentor support, and tutorials inside the workflow."),
            ("Mentors / faculty", "Need clearer coordination around availability, support, and supervised projects.", "Support mentor-linked booking and project-level visibility."),
            ("Admins", "Need ownership, inventory visibility, approvals, and usage insights.", "Provide structured admin dashboards, approval flows, notifications, and analytics."),
        ],
    )
    document.add_heading("3.3 Representative Feedback Excerpts", level=2)
    add_quote(document, '"We observed how students try to prototype and the friction they face in accessing tools and materials." [S3]')
    add_quote(document, '"Through surveys, informal interviews, and observations, we identified a recurring issue - lack of a centralized system." [S3]')
    add_quote(document, '"Feedback showed strong agreement, especially from students and admins." [S3]')
    document.add_paragraph(
        "Together, these excerpts point toward a solution that must do more than list resources. Stakeholders are asking for a system that reduces planning uncertainty, lowers coordination overhead, and makes the makerspace easier to navigate for both new and experienced users."  # noqa: E501
    )
    document.add_heading("3.4 Stakeholder-to-Feature Mapping", level=2)
    add_table(
        document,
        ["Stakeholder", "Pain or expectation", "Product response"],
        STAKEHOLDER_FEATURE_ROWS,
    )

    document.add_heading("4. Solution Refinement", level=1)
    document.add_paragraph(
        "The project did not move directly from problem discovery to the final solution. The team's Week 11 materials show a wide spread of alternatives: manual systems, lightweight digital tools, access-control mechanisms, automation-heavy shelves, and database concepts [S2]. The final solution emerged only after comparing these ideas against stakeholder needs, scope discipline, and pilot feasibility."  # noqa: E501
    )
    add_table(
        document,
        ["Version", "Core idea", "Why it changed", "Refined outcome"],
        [
            ("v1", "Manual and low-friction coordination tools such as forms, chat, and logbooks.", "Too fragmented, weak audit trail, and no live status.", "Useful as a baseline but not sufficient as a campus system."),
            ("v2", "Standalone machine-status or materials-only tools.", "Improved one layer but failed to connect discovery, booking, and governance.", "Led to a broader systems view."),
            ("v3", "Centralized information platform for tools, materials, and guidance.", "Stronger usability, but limited if approvals and support remained external.", "Extended toward integrated workflows."),
            ("v4", "Integrated Makerspace Hub with booking, training, BOM, notifications, and analytics.", "Best fit for stakeholder pain, prototype feasibility, and institutional value.", "Chosen final direction."),
        ],
    )
    document.add_heading("4.1 Concept Options and Final Homepage Direction", level=2)
    document.add_paragraph(
        "The design exploration did not stop at workflow logic. The project also generated three homepage directions aligned to the live makerspace domain: Foundry Executive, Parchment and Steel, and Vanguard Runway. The user selected Concept A, Foundry Executive, and it was later rolled across the project with adaptations for public-safe storytelling [S7][S8]."  # noqa: E501
    )
    add_table(
        document,
        ["Concept", "Core language", "Primary strength", "Primary risk", "Decision"],
        CONCEPT_COMPARISON_ROWS,
    )
    document.add_heading("4.2 PDCA Lens", level=2)
    for item in [
        "Plan: identify the campus problem space and shortlist the most actionable opportunity.",
        "Do: generate multiple solution types and map them against feasibility, cost, and user value.",
        "Check: compare alternatives against observed stakeholder pain and the need for an end-to-end workflow.",
        "Act: converge on the integrated platform and extend the prototype toward governance-heavy features such as approvals and mentor-linked booking.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_paragraph("Primary reasons for change during refinement:", style="Heading 3")
    for item in [
        "Feedback: stakeholders needed one system, not multiple disconnected fixes.",
        "Failure: manual or single-purpose tools could not hold discovery, execution, and tracking together.",
        "Cost: high-hardware automation ideas were less suitable for a rapid pilot than a software-first platform.",
        "Skills: the team could deliver a web-based prototype faster than a sensor-heavy infrastructure change.",
        "Time: an integrated but software-first prototype offered the highest value within academic constraints.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("4.3 Justified Changes from Selected Concept to Final Public Experience", level=2)
    document.add_paragraph(
        "The exemplar portfolio spends significant space on design changes between the proposed interface and the final build. The same discipline applies here. Concept A was not copied directly. It was adapted so that the public homepage could reflect the chosen visual logic without exposing internal operational dashboard content."  # noqa: E501
    )
    add_table(
        document,
        ["Aspect", "Concept A prototype", "Final implementation", "Justification"],
        CONCEPT_TO_FINAL_ROWS,
    )
    add_figure(
        document,
        assets["timeline"],
        "Figure 4. Refinement timeline from early problem scouting to the current operational prototype direction [S1][S2][S5][S6].",
    )

    document.add_heading("5. Proposed Design (The Proto Section)", level=1)
    document.add_heading("5.1 Detailed Design Brief", level=2)
    add_table(document, ["Design brief dimension", "Resolved design brief"], DESIGN_BRIEF_ROWS)
    document.add_heading("5.2 System Overview", level=2)
    document.add_paragraph(
        "The proposed prototype is an integrated makerspace operations platform for students, mentors, and admins. Students can discover machines and materials, understand training requirements, book resources, and manage project-related workflows. Mentors can expose availability and participate in guided workflows where technical support is required. Admins can manage catalog data, approvals, notifications, and analytics from the same system."  # noqa: E501
    )
    add_figure(
        document,
        assets["architecture"],
        "Figure 5. High-level system architecture for the proposed prototype, showing how user groups connect to operational modules and governance controls.",
    )
    document.add_heading("5.3 Detailed Design Specification", level=2)
    add_table(
        document,
        ["Specification area", "Required behavior", "Why it matters", "Evidence base"],
        DESIGN_SPEC_ROWS,
    )
    document.add_heading("5.4 Screen-Level Design Brief", level=2)
    add_table(
        document,
        ["Screen / surface", "Primary audience", "Purpose", "Key design decision"],
        SCREEN_BRIEF_ROWS,
    )
    document.add_heading("5.5 Subsystem Breakdown", level=2)
    document.add_heading("Discovery and Catalog", level=3)
    document.add_paragraph(
        "The catalog layer provides searchable machine and material records, category filters, specifications, safety guidance, and stock visibility. This directly addresses the documented information gap and makes remote planning possible for students before a makerspace visit."  # noqa: E501
    )
    document.add_heading("Execution and Access", level=3)
    document.add_paragraph(
        "The execution layer handles machine booking, mentor session booking, training gates, overlap detection, and mentor-supported machine workflows. This turns access into a structured, policy-aware process rather than an ad-hoc coordination effort."  # noqa: E501
    )
    document.add_heading("Project and BOM Workflows", level=3)
    document.add_paragraph(
        "The project subsystem supports team creation, mentor assignment, BOM creation, versioned material planning, and approval-driven request flows. This matters because the makerspace problem is not only about access to tools; it is also about managing project execution and material consumption with accountability."  # noqa: E501
    )
    document.add_heading("Operations, Governance, and Analytics", level=3)
    document.add_paragraph(
        "Notifications, low-stock visibility, role-based dashboards, and analytics create an operations layer for admins. Instead of reacting manually to every booking, request, or shortage, the makerspace gains a measurable operating picture."  # noqa: E501
    )
    document.add_heading("Public Discovery and Future Layer", level=3)
    document.add_paragraph(
        "The next planned increment adds public ongoing-project discovery and visibility controls, allowing the makerspace to function not only as an internal service but also as a campus engagement surface [S6]."  # noqa: E501
    )
    document.add_heading("5.6 Technical Architecture and Stack", level=2)
    add_table(
        document,
        ["Layer", "Current choice", "Technical role"],
        TECH_STACK_ROWS,
    )
    document.add_heading("5.7 Core Data Entities and Responsibilities", level=2)
    add_table(
        document,
        ["Entity", "Purpose", "Why it matters in the workflow"],
        ENTITY_ROWS,
    )
    document.add_heading("5.8 Workflow Specifications", level=2)
    document.add_paragraph(
        "The detailed technical behavior of the solution is best understood through its governed workflows. The platform is intentionally designed so that discovery, validation, execution, review, and feedback remain inside one system rather than being split across chats, forms, spreadsheets, and in-person clarification."  # noqa: E501
    )
    add_table(
        document,
        ["Workflow", "Actor sequence", "Critical gate", "Expected output"],
        WORKFLOW_ROWS,
    )
    add_figure(
        document,
        assets["workflow"],
        "Figure 6. Core operational workflow model showing how the platform moves users from visibility to validated execution and feedback.",
    )
    document.add_heading("5.9 Validation and Testing Plan", level=2)
    document.add_paragraph(
        "The exemplar report moves beyond design intent and defines how the solution should be tested. The same principle is important here. Because this prototype already exists in code, the validation plan can combine implemented evidence with pilot-stage measurement."  # noqa: E501
    )
    add_table(
        document,
        ["Validation layer", "What to check", "Current evidence", "Pilot-stage extension"],
        VALIDATION_ROWS,
    )
    document.add_paragraph("Proposed pilot metrics:", style="Heading 3")
    for item in [
        "Time taken for a student to identify a suitable machine and confirm whether they are eligible to use it.",
        "Booking completion rate without mentor/admin intervention for standard use cases.",
        "Turnaround time from BOM submission to material-request decision.",
        "Accuracy of low-stock awareness as perceived by students versus actual inventory records.",
        "Public engagement with visible projects after the public-project route is published.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("5.10 Design Rationale", level=2)
    document.add_paragraph(
        "The final design was chosen over simpler alternatives because it closes the loop between awareness, access, execution, and governance. A single booking tool would still leave materials, training, and approvals unresolved. A single inventory tool would still leave mentor support and project workflows outside the system. The integrated design is therefore the first option that matches the full shape of the documented problem."  # noqa: E501
    )
    add_figure(
        document,
        assets["coverage"],
        "Figure 7. Current prototype coverage across subsystems, derived from the live application inventory in state.md [S5].",
    )

    document.add_heading("6. Materials and Resources", level=1)
    resource_rows = [(name, description, format_inr(cost)) for name, description, cost in RESOURCE_ROWS]
    add_table(document, ["Material / Tool / Resource", "Description", "Estimated cost"], resource_rows)
    total_cost = sum(cost for _, _, cost in RESOURCE_ROWS)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Estimated pilot total: {format_inr(total_cost)}")
    run.font.bold = True
    run.font.color.rgb = ACCENT
    document.add_paragraph(
        "This estimate intentionally reflects a software-first pilot. It assumes academic or free tiers for hosting and infrastructure and allocates cost only where the campus rollout needs physical supports such as signage and onboarding collateral."  # noqa: E501
    )

    document.add_heading("7. Work Done by Team Members (Evidence-Based)", level=1)
    add_table(
        document,
        ["Member", "Primary contribution", "Evidence"],
        [(member["name"], member["role"], member["evidence"]) for member in TEAM_MEMBERS],
    )
    add_cover_note(
        document,
        "The uploaded evidence pack preserves ideation ownership clearly. The implemented prototype represents a combined team effort, but the local artifacts do not break final development work down by commit author.",
    )

    document.add_heading("8. Difficulties Faced and Learning", level=1)
    document.add_heading("8.1 Difficulties Faced", level=2)
    for item in [
        "The most feasible solution numerically was not automatically the most meaningful one, so the team had to balance scoring with strategic relevance.",
        "The makerspace problem spans visibility, access, training, material use, and governance, which made a narrow solution feel incomplete.",
        "The shared evidence pack preserved synthesized feedback more strongly than raw survey sheets, requiring careful reconstruction of the narrative from multiple artifacts.",
        "A good pilot needed to stay implementable within student constraints without collapsing into a weak, single-purpose tool.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("8.2 Technical Trade-offs and Constraints", level=2)
    for item in [
        "A software-first pilot was favored over hardware-heavy automation because it matched academic time and skill constraints better.",
        "Public pages had to be adapted away from sensitive dashboard previews while preserving the chosen design language.",
        "The current prototype is strong in operational flows, but exact pilot metrics still depend on structured user deployment and measurement.",
        "Purchase-order lifecycle work remains a clear technical gap, which affects how far material governance can go in version 1.0.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("8.3 Learning from the Process", level=2)
    for item in [
        "Problem selection must weigh institutional value and ownership, not only spreadsheet scores.",
        "Campus products fail when they digitize one step but ignore the surrounding workflow.",
        "Beginner onboarding and admin governance are product requirements, not optional extras.",
        "Analytics create institutional legitimacy because they turn usage anecdotes into visible evidence.",
        "A strong prototype is one that can already operate in pieces while still pointing clearly to version 2.0.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("9. Declaration of AI Tool(s) Usage", level=1)
    add_table(document, ["Tool", "Purpose", "Human decision applied"], AI_TOOL_ROWS)
    document.add_paragraph(
        'Example documented ideation prompt from the shared materials: "Why do students not use makerspaces despite availability?" [S3]. The final direction was not accepted automatically; it was filtered through team observation, SWOT reasoning, and feasibility checks.',
    )

    document.add_heading("10. Closing", level=1)
    document.add_heading("10.1 Conclusion", level=2)
    document.add_paragraph(
        "The Makerspace Accessibility and Information Gap is a strong ILGC problem because it is visible, actionable, and institutionally significant. The team's final answer, an integrated makerspace hub, is credible because it responds directly to stakeholder pain, survives comparison with simpler alternatives, and already exists in prototype form across multiple operational modules."  # noqa: E501
    )
    document.add_heading("10.2 Future Recommendations (Version 2.0)", level=2)
    for item in [
        "Add public ongoing-project discovery to strengthen student engagement and showcase campus making culture.",
        "Introduce QR-assisted check-in and contextual machine guidance at the point of use.",
        "Add deeper inventory intelligence such as consumption forecasting and replenishment recommendations.",
        "Extend the system to workshops, events, and purchase-order workflows once core operations stabilize.",
        "Integrate richer mentor scheduling and advanced administrative reporting for funding and planning decisions.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("10.3 References", level=2)
    for code, description in SOURCES:
        document.add_paragraph(f"{code} {description}")

    document.add_page_break()
    document.add_heading("Appendix A. Source Pack Used for This Report", level=1)
    for code, description in SOURCES:
        document.add_paragraph(f"{code} {description}", style="List Bullet")

    document.add_heading("Appendix B. Priority and SWOT Snapshot", level=1)
    add_table(
        document,
        ["Problem", "Composite score", "Interpretation"],
        [
            ("Professors repository", 15, "Highest raw feasibility score but lower strategic alignment for this team and challenge framing."),
            ("Lost-and-found management", 13, "Operationally clear and feasible, but less connected to makerspace impact."),
            ("Makerspace accessibility and information gap", 12, "Chosen because it combined clear pain, digital feasibility, and broader institutional value."),
            ("Inclusive infrastructure", 9, "High social importance but slower and more infrastructure-heavy to pilot quickly."),
        ],
    )
    add_table(
        document,
        ["SWOT quadrant", "Detailed interpretation"],
        [
            ("Strengths", "; ".join(SWOT_ITEMS["Strengths"])),
            ("Weaknesses", "; ".join(SWOT_ITEMS["Weaknesses"])),
            ("Opportunities", "; ".join(SWOT_ITEMS["Opportunities"])),
            ("Threats", "; ".join(SWOT_ITEMS["Threats"])),
        ],
    )

    document.add_heading("Appendix C. Concept Selection and Finalization Snapshot", level=1)
    add_table(
        document,
        ["Concept", "Decision summary", "How it informed the final system"],
        [
            ("Concept A - Foundry Executive", "Selected", "Provided the core operational hierarchy and role-aware command-surface logic for the final implementation [S7][S8][S9]."),
            ("Concept B - Parchment and Steel", "Not selected", "Remains useful as a reference for storytelling and documentation tone."),
            ("Concept C - Vanguard Runway", "Not selected", "Helped test the upper bound of visual energy but was not the best operational fit."),
        ],
    )
    add_table(
        document,
        ["Aspect", "Finalization logic"],
        [(row[0], row[3]) for row in CONCEPT_TO_FINAL_ROWS],
    )

    document.add_heading("Appendix D. Current Prototype Capability Snapshot", level=1)
    add_table(
        document,
        ["Subsystem", "Current status", "Evidence base"],
        [
            ("Authentication and role control", "Implemented", "state.md [S5]"),
            ("Machine and material catalogs", "Implemented", "state.md [S5]"),
            ("Booking workflows", "Implemented", "state.md [S5]"),
            ("Projects and BOMs", "Implemented", "state.md [S5]"),
            ("Material request approvals", "Implemented with minor UI gap", "state.md [S5]"),
            ("Mentor availability and training", "Implemented", "state.md [S5]"),
            ("Notifications and analytics", "Implemented", "state.md [S5]"),
            ("Purchase orders", "Placeholder / future work", "state.md [S5]"),
        ],
    )

    document.add_heading("Appendix E. Roadmap and Validation Snapshot", level=1)
    add_table(
        document,
        ["Phase", "Focus", "Deliverables", "Acceptance view"],
        ROADMAP_PHASE_ROWS,
    )
    add_table(
        document,
        ["Validation layer", "Current evidence", "Next measurement step"],
        [(row[0], row[2], row[3]) for row in VALIDATION_ROWS],
    )

    document.add_heading("Appendix F. AI Declaration Prompt Evidence", level=1)
    document.add_paragraph(
        "The uploaded ILGC presentation content explicitly references AI-assisted ideation and root-cause analysis. The preserved prompt evidence includes the following question: \"Why do students not use makerspaces despite availability?\" [S3]. This report follows that same requirement by disclosing AI usage and pairing it with the human decisions that constrained the outcome."  # noqa: E501
    )

    return save_document(document, OUTPUT_PATH)


if __name__ == "__main__":
    saved_path = build_report()
    print(f"Report generated: {saved_path}")